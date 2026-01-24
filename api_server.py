# api_server.py — drop-in replacement
# -*- coding: utf-8 -*-
from __future__ import annotations

from fastapi import FastAPI, UploadFile, File, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from pathlib import Path
import json, subprocess, sys, tempfile, shutil, os
import io
from typing import Any, Dict, Optional
from sqlalchemy.orm import Session
from wallet_api import router as wallet_router
from webhooks_razorpay import router as razorpay_webhook_router
from routes_bos_auth import router as bos_auth_router

import re

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None  # type: ignore

try:
    import docx  # python-docx
except Exception:
    docx = None  # type: ignore

try:
    import openpyxl
except Exception:
    openpyxl = None  # type: ignore

import csv

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from PIL import Image
except Exception:
    Image = None


from wallet import (
    CreditWallet,
    CreditTransaction,
    CreditPack,
    PaymentRecord,
    get_balance,
    get_or_create_wallet,
    apply_credit_topup,
    InsufficientCreditsError,
)

from db import SessionLocal
from tier_config import TIER_CONFIG
from bos_credits import charge_bos_run

app = FastAPI(title="CAIO BOS – EA API")

# Wallet + Payments routers MUST be included first
app.include_router(wallet_router)
app.include_router(razorpay_webhook_router)
app.include_router(bos_auth_router, tags=["bos-auth"])

# -------------------- Models --------------------
class EARequest(BaseModel):
    packet: dict
    user_id: int
    plan_tier: str = "demo"
    model: Optional[str] = None
    timeout_sec: int = 300
    num_predict: int = 512

class BrainRequest(BaseModel):
    packet: dict
    user_id: int
    plan_tier: str = "demo"
    model: Optional[str] = None
    timeout_sec: int = 300
    num_predict: int = 512
    brain: str

# -------------------- CORS --------------------
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# -------------------- Helpers --------------------
def repo_root() -> str:
    # Expect to run from /opt/caio-bos-validator
    return str(Path(__file__).resolve().parent)

def run_slm(
    input_json_path: str,
    brain: str,
    *,
    model: Optional[str],
    timeout_sec: int,
    num_predict: int,
) -> Dict[str, Any]:
    """
    Thin wrapper around `python -m slm.run_slm ...` so:
    - Works in dev and on Render
    - Always returns a dict (either result or structured error)
    """
    root = repo_root()
    cmd = [
        sys.executable,
        "-m",
        "slm.run_slm",
        "--input",
        input_json_path,
        "--brain",
        brain,
        "--timeout",
        str(timeout_sec),
        "--num_predict",
        str(num_predict),
    ]
    if model:
        cmd += ["--model", model]

    try:
        p = subprocess.run(
            cmd,
            cwd=root,
            capture_output=True,
            text=True,
            timeout=timeout_sec + 60,
        )
        out = {
            "stdout": p.stdout or "",
            "stderr": p.stderr or "",
            "returncode": p.returncode,
        }
        if p.returncode != 0:
            out["error"] = "SLM failed"
        # If run_slm prints JSON on stdout, prefer parsing it
        try:
            j = json.loads(p.stdout)
            if isinstance(j, dict):
                return j
        except Exception:
            pass
        return out
    except Exception as e:
        return {
            "error": "SLM failed",
            "stdout": "",
            "stderr": str(e),
        }

def _extract_text_from_upload(filename: str, data: bytes) -> str:
    """
    Extract usable text from common file types with fallbacks.
    - PDFs: pypdf -> pdfplumber -> OCR (optional) if still too short
    - DOCX: paragraphs + tables
    - XLSX: all sheets, row-wise
    - CSV/TSV: delimiter sniff + robust decoding
    - Images: OCR (optional)
    Always returns a string; never raises UnicodeDecodeError.
    """
    name = (filename or "").lower().strip()

    # --- Safety caps ---
    MAX_TEXT_CHARS = 250_000      # hard cap to keep packets sane
    PDF_MIN_TEXT_CHARS = 3000     # below this, try fallback (plumber/OCR)
    MAX_PDF_PAGES_OCR = 12        # OCR can be expensive; cap pages
    MAX_XLSX_ROWS_PER_SHEET = 2000
    MAX_XLSX_CELLS_PER_ROW = 50

    def _cap(s: str) -> str:
        s = (s or "").strip()
        return s[:MAX_TEXT_CHARS]

    def _decode_bytes(b: bytes) -> str:
        # Try common encodings. utf-8-sig handles BOM.
        for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
            try:
                return b.decode(enc)
            except Exception:
                continue
        return b.decode("utf-8", errors="ignore")

    def _is_probably_binary(b: bytes) -> bool:
        # Heuristic: lots of NUL bytes suggests binary
        return b.count(b"\x00") > 10

    def _extract_pdf_pypdf(b: bytes) -> str:
        if PdfReader is None:
            return ""
        try:
            reader = PdfReader(io.BytesIO(b))
            parts = []
            for page in reader.pages:
                t = page.extract_text() or ""
                t = t.strip()
                if t:
                    parts.append(t)
            return "\n\n".join(parts).strip()
        except Exception:
            return ""

    def _extract_pdf_pdfplumber(b: bytes) -> str:
        if pdfplumber is None:
            return ""
        try:
            parts = []
            with pdfplumber.open(io.BytesIO(b)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    t = t.strip()
                    if t:
                        parts.append(t)
            return "\n\n".join(parts).strip()
        except Exception:
            return ""

    def _extract_image_ocr(img_bytes: bytes) -> str:
        if pytesseract is None or Image is None:
            return ""
        try:
            img = Image.open(io.BytesIO(img_bytes))
            # Basic OCR; keep it simple/reliable
            txt = pytesseract.image_to_string(img)
            return (txt or "").strip()
        except Exception:
            return ""

    def _extract_pdf_ocr(b: bytes) -> str:
        """
        OCR PDF by rendering pages to images via pdftoppm (poppler-utils)
        and running pytesseract. Uses a cap on pages.
        """
        if pytesseract is None or Image is None:
            return ""
        # If poppler isn't installed, this will fail; we catch below.
        import tempfile
        import subprocess
        import os
        try:
            with tempfile.TemporaryDirectory() as td:
                pdf_path = os.path.join(td, "in.pdf")
                with open(pdf_path, "wb") as f:
                    f.write(b)

                # Render to PNGs: page-1.png, page-2.png, ...
                # -f 1 -l N caps pages
                out_prefix = os.path.join(td, "page")
                cmd = ["pdftoppm", "-png", "-f", "1", "-l", str(MAX_PDF_PAGES_OCR), pdf_path, out_prefix]
                subprocess.run(cmd, check=True, capture_output=True)

                parts = []
                # Collect rendered pages in order
                for i in range(1, MAX_PDF_PAGES_OCR + 1):
                    img_path = f"{out_prefix}-{i}.png"
                    if not os.path.exists(img_path):
                        break
                    with open(img_path, "rb") as imf:
                        img_bytes = imf.read()
                    t = _extract_image_ocr(img_bytes)
                    if t:
                        parts.append(f"[OCR Page {i}]\n{t}")
                return "\n\n".join(parts).strip()
        except Exception:
            return ""

    # -------------------------
    # PDF
    # -------------------------
    if name.endswith(".pdf"):
        # 1) pypdf
        t1 = _extract_pdf_pypdf(data)
        best = t1

        # 2) pdfplumber fallback
        if len(best) < PDF_MIN_TEXT_CHARS:
            t2 = _extract_pdf_pdfplumber(data)
            if len(t2) > len(best):
                best = t2

        # 3) OCR fallback (optional)
        if len(best) < PDF_MIN_TEXT_CHARS:
            t3 = _extract_pdf_ocr(data)
            if len(t3) > len(best):
                best = t3

        return _cap(best)

    # -------------------------
    # DOCX
    # -------------------------
    if name.endswith(".docx"):
        if docx is None:
            return ""
        try:
            d = docx.Document(io.BytesIO(data))
            parts = []

            # paragraphs
            for p in d.paragraphs:
                txt = (p.text or "").strip()
                if txt:
                    parts.append(txt)

            # tables
            for table in d.tables:
                for row in table.rows:
                    cells = []
                    for cell in row.cells:
                        ct = (cell.text or "").strip()
                        if ct:
                            cells.append(ct)
                    if cells:
                        parts.append(" | ".join(cells))

            return _cap("\n".join(parts))
        except Exception:
            return ""

    # -------------------------
    # Excel
    # -------------------------
    if name.endswith(".xlsx"):
        if openpyxl is None:
            return ""
        try:
            wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
            parts = []
            for sheet in wb.worksheets:
                parts.append(f"[Sheet: {sheet.title}]")
                row_count = 0
                for row in sheet.iter_rows(values_only=True):
                    if row_count >= MAX_XLSX_ROWS_PER_SHEET:
                        parts.append("[TRUNCATED: too many rows]")
                        break
                    row_count += 1

                    cells = []
                    for c in row[:MAX_XLSX_CELLS_PER_ROW]:
                        if c is None:
                            continue
                        s = str(c).strip()
                        if s:
                            cells.append(s)
                    if cells:
                        parts.append(" | ".join(cells))
            return _cap("\n".join(parts))
        except Exception:
            return ""

    # -------------------------
    # CSV / TSV
    # -------------------------
    if name.endswith((".csv", ".tsv")):
        try:
            text = _decode_bytes(data)
            # Sniff delimiter (fallback to comma/tsv)
            sample = text[:4096]
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=[",", "\t", ";", "|"])
                delim = dialect.delimiter
            except Exception:
                delim = "\t" if name.endswith(".tsv") else ","

            reader = csv.reader(io.StringIO(text), delimiter=delim)
            out_lines = []
            for i, row in enumerate(reader):
                if i > 3000:
                    out_lines.append("[TRUNCATED: too many rows]")
                    break
                # Keep rows bounded
                row = [cell.strip() for cell in row[:60] if cell and cell.strip()]
                if row:
                    out_lines.append(" | ".join(row))
            return _cap("\n".join(out_lines))
        except Exception:
            # Last resort decode
            return _cap(_decode_bytes(data))

    # -------------------------
    # Plain text-like
    # -------------------------
    if name.endswith((".txt", ".md", ".json", ".yaml", ".yml", ".log")):
        return _cap(_decode_bytes(data))

    # -------------------------
    # Images (OCR)
    # -------------------------
    if name.endswith((".png", ".jpg", ".jpeg", ".webp")):
        t = _extract_image_ocr(data)
        return _cap(t)

    # -------------------------
    # Fallback
    # -------------------------
    if _is_probably_binary(data):
        # Don't attempt to "decode" arbitrary binaries meaningfully
        return ""
    return _cap(_decode_bytes(data))

def _extract_text_with_meta(filename: str, data: bytes) -> tuple[str, dict]:
    """
    Returns (text, meta). Meta includes:
    - methods_tried
    - chosen_method
    - text_len
    - quality_flags
    - hints (e.g., quote-like filename but no money signals)
    """

    name = (filename or "").lower().strip()

    # ---- Tuning knobs ----
    MAX_TEXT_CHARS = 250_000
    PDF_MIN_TEXT_CHARS = 3000

    OCR_DPI = 300
    OCR_HEAD_PAGES = 4
    OCR_TAIL_PAGES = 4
    OCR_MAX_TOTAL_PAGES = 12  # safety

    MAX_XLSX_ROWS_PER_SHEET = 2000
    MAX_XLSX_CELLS_PER_ROW = 50
    MAX_CSV_ROWS = 3000
    MAX_CSV_COLS = 60

    meta = {
        "filename": filename,
        "ext": name.split(".")[-1] if "." in name else "",
        "methods_tried": [],
        "chosen_method": None,
        "text_len": 0,
        "quality_flags": [],
        "hints": {},
    }

    def cap(s: str) -> str:
        s = (s or "").strip()
        if len(s) > MAX_TEXT_CHARS:
            return s[:MAX_TEXT_CHARS]
        return s

    def decode_bytes(b: bytes) -> str:
        for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
            try:
                return b.decode(enc)
            except Exception:
                continue
        return b.decode("utf-8", errors="ignore")

    def is_probably_binary(b: bytes) -> bool:
        return b.count(b"\x00") > 10

    def money_signals(s: str) -> bool:
        if not s:
            return False
        s2 = s.lower()
        patterns = [
            r"₹", r"\$", r"\binr\b", r"\busd\b", r"\brs\.?\b", r"\bgst\b", r"\btax\b",
            r"\btotal\b", r"\bsubtotal\b", r"\bgrand total\b", r"\bamount\b",
            r"\bquotation\b", r"\binvoice\b", r"\bpricing\b",
            r"\b\d{1,3}(,\d{3})+(\.\d+)?\b",  # 1,23,456 style
            r"\b\d+\.\d+\b",
        ]
        return any(re.search(p, s2) for p in patterns)

    def quote_like_filename(n: str) -> bool:
        return bool(re.search(r"(quote|quotation|invoice|pricing|estimate|proposal)", n.lower()))

    # --- PDF extractors ---
    def pdf_pypdf(b: bytes) -> str:
        meta["methods_tried"].append("pdf:pypdf")
        if PdfReader is None:
            return ""
        try:
            reader = PdfReader(io.BytesIO(b))
            parts = []
            for page in reader.pages:
                t = (page.extract_text() or "").strip()
                if t:
                    parts.append(t)
            return "\n\n".join(parts).strip()
        except Exception:
            return ""

    def pdf_plumber(b: bytes) -> str:
        meta["methods_tried"].append("pdf:pdfplumber")
        if pdfplumber is None:
            return ""
        try:
            parts = []
            with pdfplumber.open(io.BytesIO(b)) as pdf:
                for page in pdf.pages:
                    t = (page.extract_text() or "").strip()
                    if t:
                        parts.append(t)
            return "\n\n".join(parts).strip()
        except Exception:
            return ""

    def image_ocr(img_bytes: bytes) -> str:
        if pytesseract is None or Image is None:
            return ""
        try:
            img = Image.open(io.BytesIO(img_bytes))
            txt = pytesseract.image_to_string(img)
            return (txt or "").strip()
        except Exception:
            return ""

    def pdf_ocr_first_last(b: bytes) -> str:
        """
        OCR first N pages and last N pages.
        Uses pdftoppm (poppler-utils) -> PNG -> pytesseract.
        """
        meta["methods_tried"].append("pdf:ocr_first_last")
        if pytesseract is None or Image is None:
            return ""

        import tempfile, subprocess, os

        try:
            with tempfile.TemporaryDirectory() as td:
                pdf_path = os.path.join(td, "in.pdf")
                with open(pdf_path, "wb") as f:
                    f.write(b)

                # Determine page count (best-effort)
                # If PdfReader exists, use it to count pages; else assume OCR_HEAD+OCR_TAIL.
                page_count = None
                try:
                    if PdfReader is not None:
                        reader = PdfReader(io.BytesIO(b))
                        page_count = len(reader.pages)
                except Exception:
                    page_count = None

                head_n = OCR_HEAD_PAGES
                tail_n = OCR_TAIL_PAGES

                if page_count is not None:
                    # cap total OCR pages
                    if head_n + tail_n > OCR_MAX_TOTAL_PAGES:
                        head_n = min(head_n, OCR_MAX_TOTAL_PAGES)
                        tail_n = max(0, OCR_MAX_TOTAL_PAGES - head_n)
                    tail_start = max(1, page_count - tail_n + 1)
                else:
                    # unknown count: just OCR first OCR_MAX_TOTAL_PAGES pages
                    head_n = min(head_n + tail_n, OCR_MAX_TOTAL_PAGES)
                    tail_n = 0
                    tail_start = None

                parts = []

                # OCR head pages
                if head_n > 0:
                    out_prefix = os.path.join(td, "head")
                    cmd = ["pdftoppm", "-r", str(OCR_DPI), "-png", "-f", "1", "-l", str(head_n), pdf_path, out_prefix]
                    subprocess.run(cmd, check=True, capture_output=True)

                    for i in range(1, head_n + 1):
                        img_path = f"{out_prefix}-{i}.png"
                        if not os.path.exists(img_path):
                            break
                        with open(img_path, "rb") as imf:
                            t = image_ocr(imf.read())
                        if t:
                            parts.append(f"[OCR Head Page {i}]\n{t}")

                # OCR tail pages
                if tail_n > 0 and tail_start is not None:
                    out_prefix = os.path.join(td, "tail")
                    cmd = ["pdftoppm", "-r", str(OCR_DPI), "-png", "-f", str(tail_start), "-l", str(page_count), pdf_path, out_prefix]
                    subprocess.run(cmd, check=True, capture_output=True)

                    # tail images start at 1 in output naming
                    for i in range(1, tail_n + 1):
                        img_path = f"{out_prefix}-{i}.png"
                        if not os.path.exists(img_path):
                            break
                        with open(img_path, "rb") as imf:
                            t = image_ocr(imf.read())
                        if t:
                            parts.append(f"[OCR Tail Page {tail_start + i - 1}]\n{t}")

                return "\n\n".join(parts).strip()

        except Exception:
            return ""

    # -------------------------
    # PDF
    # -------------------------
    if name.endswith(".pdf"):
        t1 = pdf_pypdf(data)
        best = t1
        best_method = "pdf:pypdf"

        if len(best) < PDF_MIN_TEXT_CHARS:
            t2 = pdf_plumber(data)
            if len(t2) > len(best):
                best = t2
                best_method = "pdf:pdfplumber"

        if len(best) < PDF_MIN_TEXT_CHARS:
            t3 = pdf_ocr_first_last(data)
            if len(t3) > len(best):
                best = t3
                best_method = "pdf:ocr_first_last"

        best = cap(best)
        meta["chosen_method"] = best_method
        meta["text_len"] = len(best)

        # Quality flags + UX hints (Option A + C)
        if meta["text_len"] < PDF_MIN_TEXT_CHARS:
            meta["quality_flags"].append("LOW_TEXT_PDF")

        if quote_like_filename(filename):
            meta["hints"]["quote_like_filename"] = True
            if not money_signals(best):
                meta["quality_flags"].append("LIKELY_QUOTE_PRICING_NOT_EXTRACTED")

        return best, meta

    # -------------------------
    # DOCX (paragraphs + tables)
    # -------------------------
    if name.endswith(".docx"):
        meta["methods_tried"].append("docx:python-docx")
        if docx is None:
            meta["quality_flags"].append("DOCX_PARSER_MISSING")
            return "", meta
        try:
            d = docx.Document(io.BytesIO(data))
            parts = []

            for p in d.paragraphs:
                txt = (p.text or "").strip()
                if txt:
                    parts.append(txt)

            for table in d.tables:
                for row in table.rows:
                    cells = []
                    for cell in row.cells:
                        ct = (cell.text or "").strip()
                        if ct:
                            cells.append(ct)
                    if cells:
                        parts.append(" | ".join(cells))

            out = cap("\n".join(parts))
            meta["chosen_method"] = "docx:python-docx"
            meta["text_len"] = len(out)
            if meta["text_len"] < 200:
                meta["quality_flags"].append("LOW_TEXT_DOCX")
            return out, meta
        except Exception:
            meta["quality_flags"].append("DOCX_EXTRACT_FAILED")
            return "", meta

    # -------------------------
    # XLSX
    # -------------------------
    if name.endswith(".xlsx"):
        meta["methods_tried"].append("xlsx:openpyxl")
        if openpyxl is None:
            meta["quality_flags"].append("XLSX_PARSER_MISSING")
            return "", meta
        try:
            wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
            parts = []
            for sheet in wb.worksheets:
                parts.append(f"[Sheet: {sheet.title}]")
                row_count = 0
                for row in sheet.iter_rows(values_only=True):
                    if row_count >= MAX_XLSX_ROWS_PER_SHEET:
                        parts.append("[TRUNCATED: too many rows]")
                        break
                    row_count += 1

                    cells = []
                    for c in row[:MAX_XLSX_CELLS_PER_ROW]:
                        if c is None:
                            continue
                        s = str(c).strip()
                        if s:
                            cells.append(s)
                    if cells:
                        parts.append(" | ".join(cells))

            out = cap("\n".join(parts))
            meta["chosen_method"] = "xlsx:openpyxl"
            meta["text_len"] = len(out)
            if meta["text_len"] < 200:
                meta["quality_flags"].append("LOW_TEXT_XLSX")
            return out, meta
        except Exception:
            meta["quality_flags"].append("XLSX_EXTRACT_FAILED")
            return "", meta

    # -------------------------
    # CSV / TSV
    # -------------------------
    if name.endswith((".csv", ".tsv")):
        meta["methods_tried"].append("csv:sniff")
        try:
            text = decode_bytes(data)
            sample = text[:4096]
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=[",", "\t", ";", "|"])
                delim = dialect.delimiter
            except Exception:
                delim = "\t" if name.endswith(".tsv") else ","

            rdr = csv.reader(io.StringIO(text), delimiter=delim)
            out_lines = []
            for i, row in enumerate(rdr):
                if i >= MAX_CSV_ROWS:
                    out_lines.append("[TRUNCATED: too many rows]")
                    break
                row = [(c or "").strip() for c in row[:MAX_CSV_COLS]]
                row = [c for c in row if c]
                if row:
                    out_lines.append(" | ".join(row))

            out = cap("\n".join(out_lines))
            meta["chosen_method"] = f"csv:{'tsv' if delim=='\\t' else 'delim'}"
            meta["text_len"] = len(out)
            if meta["text_len"] < 200:
                meta["quality_flags"].append("LOW_TEXT_CSV")
            return out, meta
        except Exception:
            out = cap(decode_bytes(data))
            meta["chosen_method"] = "csv:decode_fallback"
            meta["text_len"] = len(out)
            return out, meta

    # -------------------------
    # Plain text-like
    # -------------------------
    if name.endswith((".txt", ".md", ".json", ".yaml", ".yml", ".log")):
        meta["methods_tried"].append("text:decode")
        out = cap(decode_bytes(data))
        meta["chosen_method"] = "text:decode"
        meta["text_len"] = len(out)
        return out, meta

    # -------------------------
    # Images (OCR)
    # -------------------------
    if name.endswith((".png", ".jpg", ".jpeg", ".webp")):
        meta["methods_tried"].append("img:ocr")
        out = cap(image_ocr(data))
        meta["chosen_method"] = "img:ocr"
        meta["text_len"] = len(out)
        if meta["text_len"] < 50:
            meta["quality_flags"].append("LOW_TEXT_IMAGE_OCR")
        return out, meta

    # -------------------------
    # Fallback
    # -------------------------
    if is_probably_binary(data):
        meta["methods_tried"].append("fallback:binary")
        meta["chosen_method"] = "fallback:binary"
        meta["text_len"] = 0
        meta["quality_flags"].append("BINARY_UNSUPPORTED")
        return "", meta

    meta["methods_tried"].append("fallback:decode")
    out = cap(decode_bytes(data))
    meta["chosen_method"] = "fallback:decode"
    meta["text_len"] = len(out)
    return out, meta


# -------------------- Routes --------------------
@app.get("/")
def root():
    return {"ok": True, "service": "caio-bos"}

@app.get("/health")
def health():
    return {"ok": True}

@app.get("/welcome")
def welcome():
    return {"ok": True, "message": "Welcome to CAIO BOS"}

@app.post("/run-ea")
def run_ea(payload: EARequest):
    # --- Guard: prevent empty Decision Review packets (avoid timeouts / fluff) ---
    pkt = payload.packet or {}
    findings = pkt.get("findings") or []
    insights_map = pkt.get("insights") or {}
    document_text = (pkt.get("document_text") or pkt.get("text") or "").strip()

    has_insights = False
    if isinstance(insights_map, dict):
        for b in ["cfo", "cmo", "coo", "chro", "cpo", "ea"]:
            if insights_map.get(b):
                has_insights = True
                break
    else:
        has_insights = bool(insights_map)

    if (not findings) and (not has_insights) and (not document_text):
        raise HTTPException(
            status_code=400,
            detail="Decision Review requires findings/insights or document_text. Upload a file (Analyze) or provide a populated validator packet."
        )

    # Charge credits (if configured)
    try:
        charge_bos_run(payload.user_id, payload.plan_tier)
    except InsufficientCreditsError as e:
        raise HTTPException(status_code=402, detail=str(e))
    except Exception:
        # If charging fails, still allow run; you can tighten later
        pass

    # Save packet to temp file
    with tempfile.NamedTemporaryFile(suffix=".json", delete=False, mode="w", encoding="utf-8") as tf:
        json.dump(payload.packet, tf, ensure_ascii=False, indent=2)
        tmp_in = tf.name

    out = run_slm(
        tmp_in,
        "ea",
        model=payload.model,
        timeout_sec=payload.timeout_sec,
        num_predict=payload.num_predict,
    )
    if "error" in out and "ui" not in out:
        return {"ui": out}
    return out


@app.post("/run-brain")
def run_brain(payload: BrainRequest):
    # Charge credits (if configured)
    try:
        charge_bos_run(payload.user_id, payload.plan_tier)
    except InsufficientCreditsError as e:
        raise HTTPException(status_code=402, detail=str(e))
    except Exception:
        pass

    with tempfile.NamedTemporaryFile(suffix=".json", delete=False, mode="w", encoding="utf-8") as tf:
        json.dump(payload.packet, tf, ensure_ascii=False, indent=2)
        tmp_in = tf.name

    out = run_slm(
        tmp_in,
        payload.brain,
        model=payload.model,
        timeout_sec=payload.timeout_sec,
        num_predict=payload.num_predict,
    )
    if "error" in out and "ui" not in out:
        return {"ui": out}
    return out

@app.post("/upload-and-ea")
async def upload_and_ea(
    file: UploadFile = File(...),
    timeout_sec: int = 300,
    num_predict: int = 512,
    model: Optional[str] = None,
):
    """
    Upload a file and run EA.

    - If file is JSON: treat as BOS/validator packet JSON (existing behavior).
    - Else: extract readable text and wrap it into a packet JSON for EA.

    Returns: {"ui": ...} compatible with BOSSummary.
    """
    filename = file.filename or "upload"
    raw = await file.read()

    # JSON packet path (backward compatible)
    if filename.lower().endswith(".json"):
        with tempfile.NamedTemporaryFile(suffix=".json", delete=False, mode="wb") as tf:
            tf.write(raw)
            tmp_in = tf.name
        out = run_slm(
            tmp_in,
            "ea",
            model=model,
            timeout_sec=timeout_sec,
            num_predict=num_predict,
        )
        if "error" in out and "ui" not in out:
            return {"ui": out}
        return {"ui": out.get("ui") or out}

    # Document path (PDF/DOCX/TXT/other)
    text, extract_meta = _extract_text_with_meta(filename, raw)
    
    print(
    f"[UPLOAD] filename={filename} "
    f"len={len(text)} "
    f"preview={text[:200]!r}"
)

    print(f"[EXTRACT] chosen={extract_meta.get('chosen_method')} flags={extract_meta.get('quality_flags')}")

    if not text or len(text.strip()) < 20:
        return {
            "ui": {
                "error": "No readable text extracted from upload",
                "stdout": "",
                "stderr": "",
            }
        }

    # Wrap extracted text into a packet JSON for EA
    packet: Dict[str, Any] = {
        "label": "Uploaded Document",
        "source": {
            "filename": filename,
            "content_type": file.content_type,
            "size_bytes": len(raw),
        },
        "document_text": text[:200000],  # safety cap
        "facts": {},
        "meta": {"ingest": "upload-and-ea"},
    }
    
    packet["meta"]["doc_text_len"] = len(text)
    packet["meta"]["doc_text_preview"] = text[:400]

    with tempfile.NamedTemporaryFile(suffix=".json", delete=False, mode="w", encoding="utf-8") as tf:
        json.dump(packet, tf, ensure_ascii=False, indent=2)
        tmp_in = tf.name

    out = run_slm(
        tmp_in,
        "ea",
        model=model,
        timeout_sec=timeout_sec,
        num_predict=num_predict,
    )
    
    # Attach extraction metadata
    packet["meta"]["extract"] = extract_meta
    
    # Build warnings for UI
    warnings = []
    flags = extract_meta.get("quality_flags") or []
    if "LIKELY_QUOTE_PRICING_NOT_EXTRACTED" in flags:
        warnings.append(
            "Pricing/quotation terms may be embedded as an image/table and were not extracted reliably. "
            "Upload the quotation as XLSX/CSV or a text-based PDF, or upload the quotation pages separately."
        )
    elif "LOW_TEXT_PDF" in flags:
        warnings.append(
            "This PDF contains limited extractable text (possibly scanned or table-heavy). "
            "Results may be incomplete; consider uploading a text-based PDF or an XLSX/CSV version."
        )
    
    packet["meta"]["warnings"] = warnings
    
    if "error" in out and "ui" not in out:
        return {"ui": out}
    ui_obj = out.get("ui") or out
    if isinstance(ui_obj, dict):
        ui_obj.setdefault("warnings", [])
        ui_obj["warnings"].extend(packet["meta"].get("warnings", []))
        ui_obj["extract_meta"] = extract_meta
    return {"ui": ui_obj}

    
    
