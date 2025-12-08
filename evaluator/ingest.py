# -*- coding: utf-8 -*-
from typing import Dict, Any
from pathlib import Path
import json
import pandas as pd

# Optional imports for rich formats
try:
    from docx import Document  # python-docx
except Exception:
    Document = None

try:
    from pypdf import PdfReader  # pypdf (lightweight)
except Exception:
    PdfReader = None

BRAINS = ["cfo", "coo", "cmo", "chro", "cpo"]

def _df_from_csv(p: Path) -> pd.DataFrame:
    return pd.read_csv(p)

def _df_from_excel_first(p: Path) -> pd.DataFrame:
    return pd.read_excel(p)

def _dfs_from_excel_all(p: Path) -> Dict[str, pd.DataFrame]:
    x = pd.ExcelFile(p)
    return {name.lower(): x.parse(name) for name in x.sheet_names}

def _df_from_json(p: Path) -> pd.DataFrame:
    obj = json.loads(p.read_text(encoding="utf-8"))

    # Case 1: already a list of rows
    if isinstance(obj, list):
        return pd.DataFrame(obj)

    # Case 2: dict -> try known tabular keys first, then any list-of-dicts
    if isinstance(obj, dict):
        preferred = [
            "pnl", "balance_sheet", "cashflow", "channel_report",
            "orders", "hr", "inventory", "table", "data"
        ]
        for k in preferred:
            v = obj.get(k)
            if isinstance(v, list) and v and isinstance(v[0], dict):
                return pd.DataFrame(v)

        for v in obj.values():
            if isinstance(v, list) and v and isinstance(v[0], dict):
                return pd.DataFrame(v)

        # Fallback: flatten dict into single row (lets resolver work with text/meta)
        return pd.json_normalize(obj, max_level=1)

    raise ValueError("Unsupported JSON structure")


def _tables_from_docx(p: Path):
    if Document is None:
        return []
    doc = Document(str(p))
    tables = []
    for t in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in t.rows]
        if rows and any(any(c for c in r) for r in rows):
            df = pd.DataFrame(rows)
            df.columns = df.iloc[0].fillna("").astype(str)
            df = df[1:].reset_index(drop=True)
            tables.append(df)
    return tables

def _text_from_docx(p: Path) -> str:
    if Document is None:
        return ""
    doc = Document(str(p))
    return "\n".join([para.text for para in doc.paragraphs]).strip()

def _text_from_pdf(p: Path) -> str:
    if PdfReader is None:
        return ""
    reader = PdfReader(str(p))
    chunks = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            chunks.append("")
    return "\n".join(chunks).strip()

def _text_from_txt(p: Path) -> str:
    return p.read_text(encoding="utf-8", errors="ignore")

def _unstructured_df(text: str) -> pd.DataFrame:
    # minimal schema so the engine can still run
    return pd.DataFrame({"text": [text]})

def _route_excel_sheets_to_brains(sheets: Dict[str, pd.DataFrame]) -> Dict[str, pd.DataFrame]:
    # 1: direct sheet name hit (cfo/cmo/...)
    out = {}
    for b in BRAINS:
        if b in sheets:
            out[b] = sheets[b]
    # 2: fuzzy mapping
    hints = {
        "cfo": ["finance", "fin", "p&l", "pnl", "pl", "bs", "balance", "cash"],
        "cmo": ["marketing", "mkt", "channel", "utm", "ads", "campaign"],
        "coo": ["ops", "operations", "supply", "throughput", "production"],
        "chro": ["hr", "people", "org", "headcount", "attrition"],
        "cpo": ["hiring", "recruit", "talent", "offers", "ctc", "salary"],
    }
    for b in BRAINS:
        if b in out: continue
        for h in hints[b]:
            match = next((df for name, df in sheets.items() if h in name), None)
            if match is not None:
                out[b] = match; break
    # 3: fallback to first sheet if still missing
    first_df = next(iter(sheets.values())) if sheets else pd.DataFrame()
    for b in BRAINS:
        out.setdefault(b, first_df)
    return out

def load_any(input_path: str) -> Dict[str, pd.DataFrame]:
    """
    Returns {brain: DataFrame}. Handles: .xlsx/.xls, .csv, .json, .docx, .pdf, .txt
    """
    p = Path(input_path)
    ext = p.suffix.lower()

    if ext in (".xlsx", ".xls"):
        sheets = _dfs_from_excel_all(p)
        return _route_excel_sheets_to_brains(sheets)

    if ext == ".csv":
        df = _df_from_csv(p)
        return {b: df for b in BRAINS}

    if ext == ".json":
        df = _df_from_json(p)
        return {b: df for b in BRAINS}

    if ext == ".docx":
        # Prefer tables if present; else text
        tables = _tables_from_docx(p)
        if tables:
            # Use largest table as a generic source for all brains (intent resolver will adapt)
            biggest = max(tables, key=lambda d: d.shape[0] * d.shape[1])
            return {b: biggest for b in BRAINS}
        text = _text_from_docx(p)
        df = _unstructured_df(text)
        return {b: df for b in BRAINS}

    if ext == ".pdf":
        text = _text_from_pdf(p)
        df = _unstructured_df(text)
        return {b: df for b in BRAINS}

    if ext == ".txt":
        text = _text_from_txt(p)
        df = _unstructured_df(text)
        return {b: df for b in BRAINS}

    raise ValueError(f"Unsupported input file type: {ext}")
